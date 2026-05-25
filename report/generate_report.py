# -*- coding: utf-8 -*-
"""Sinh báo cáo đánh giá dự án FishShop theo Debai.txt."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullets(doc: Document, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_status_table(doc: Document, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Hạng mục"
    hdr[1].text = "Trạng thái"
    hdr[2].text = "Ghi chú"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(c, "1F3A5F")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (item, status, note) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = item
        row[1].text = status
        row[2].text = note


doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── TRANG TIÊU ĐỀ ─────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BÁO CÁO ĐÁNH GIÁ ĐỒ ÁN")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Chủ đề 4 – Phát triển ứng dụng Microservice")
sub_run.italic = True
sub_run.font.size = Pt(14)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_run = sub2.add_run("Dự án: FishShop – Hệ thống bán cá cảnh trực tuyến")
sub2_run.font.size = Pt(13)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run("Báo cáo đối chiếu yêu cầu đề bài, hiện trạng đã làm được,\ncác vấn đề còn tồn tại và phương án xử lý.")
info_run.italic = True

doc.add_paragraph()

# ── 1. GIỚI THIỆU ─────────────────────────────────────────
add_heading(doc, "1. Giới thiệu chung", level=1)
add_para(
    doc,
    "Dự án FishShop là một ứng dụng thương mại điện tử bán cá cảnh và phụ kiện, "
    "được xây dựng theo kiến trúc microservice nhằm đáp ứng yêu cầu của Chủ đề 4 "
    "(Phát triển ứng dụng Microservice). Backend sử dụng .NET 8, giao tiếp đồng bộ "
    "qua RESTful API và giao tiếp bất đồng bộ qua message queue (RabbitMQ). "
    "Frontend gồm hai ứng dụng React/Vite riêng biệt: trang quản trị và trang khách hàng."
)
add_para(doc, "Mục tiêu của báo cáo này là:", bold=True)
add_bullets(doc, [
    "Đối chiếu hiện trạng dự án với từng yêu cầu trong file Debai.txt.",
    "Liệt kê các hạng mục đã hoàn thành kèm bằng chứng trong mã nguồn.",
    "Chỉ ra các vấn đề còn tồn tại, giải thích bản chất của từng vấn đề.",
    "Đề xuất phương án và phương pháp cụ thể để khắc phục.",
])

# ── 2. KIẾN TRÚC TỔNG QUAN ────────────────────────────────
add_heading(doc, "2. Kiến trúc tổng quan đã triển khai", level=1)
add_para(doc, "Hệ thống hiện tại gồm các thành phần chính sau:")
add_bullets(doc, [
    "4 microservice độc lập: API.User, API.Product, API.Order, API.Content.",
    "API Gateway (FishShop.Gateway) dùng YARP làm điểm vào duy nhất từ frontend.",
    "Message broker RabbitMQ phục vụ luồng bất đồng bộ giữa Order và Product.",
    "4 cơ sở dữ liệu SQL Server tách biệt: FishShop_User, FishShop_Product, FishShop_Order, FishShop_Content.",
    "Centralized logging với Seq + Serilog.",
    "Frontend Admin (FE/) và Frontend Customer (FE-Customer/) gọi API qua Gateway.",
    "Docker Compose cho RabbitMQ, Seq, 4 service và Gateway.",
])

# ── 3. NHỮNG GÌ ĐÃ LÀM ĐƯỢC ───────────────────────────────
add_heading(doc, "3. Những hạng mục đã hoàn thành", level=1)
add_para(
    doc,
    "Bảng dưới đây đối chiếu từng yêu cầu bắt buộc trong Debai.txt với hiện trạng dự án:"
)

add_status_table(doc, [
    ("3.1 Có ít nhất 3 microservice",
     "Đạt",
     "4 service: User, Product, Order, Content."),
    ("3.2 RESTful API đúng tài nguyên & HTTP method",
     "Đạt",
     "Route theo dạng /api/{service}/{resource}, CRUD chuẩn (GET/POST/PUT/PATCH/DELETE)."),
    ("3.3 Sử dụng message queue trong ≥ 1 luồng nghiệp vụ",
     "Đạt",
     "Order publish OrderCreatedEvent, Product consume để trừ kho qua queue order-created-queue."),
    ("3.4 Mỗi service có cơ sở dữ liệu riêng",
     "Đạt",
     "4 DB SQL Server tách biệt, mỗi service chỉ kết nối DB của mình."),
    ("3.5 Xử lý lỗi cơ bản (retry, logging)",
     "Đạt cơ bản",
     "MassTransit UseMessageRetry(3 lần, 5s); GlobalExceptionFilter cho REST; Serilog ghi log lỗi."),
    ("4. API Gateway",
     "Đạt",
     "YARP làm reverse proxy + middleware kiểm tra JWT và chèn X-Api-Key."),
    ("4. Xác thực JWT",
     "Đạt",
     "Đăng nhập trả về JWT; Gateway validate token; các endpoint cần đăng nhập dùng Bearer."),
    ("4. Triển khai bằng Docker / Docker Compose",
     "Đạt một phần",
     "Có compose cho RabbitMQ, Seq, 4 service, Gateway. SQL Server hiện chạy ngoài compose."),
    ("4. Centralized logging",
     "Đạt",
     "Serilog ghi ra Console + File + Seq tại http://localhost:5341."),
    ("4. Có log theo dõi request & message",
     "Đạt",
     "UseSerilogRequestLogging ở Gateway; consumer log mỗi message nhận được."),
])

add_para(doc, "Một số điểm nổi bật trong hiện trạng:", bold=True)
add_bullets(doc, [
    "Gateway chèn header X-Api-Key vào request nội bộ; các service kiểm tra header này để bảo vệ khỏi truy cập trực tiếp (thuộc tính ApiKeyAttribute trong netcore.Commons).",
    "Order Service phát event OrderCreatedEvent sau khi commit transaction tạo đơn; Product Service nhận event và gọi InventoryService.ExportStockAsync để trừ kho – đúng tinh thần loose coupling.",
    "Có cấu hình retry 3 lần (mỗi lần cách 5 giây) cho consumer kho.",
    "Có tài liệu hướng dẫn chạy đầy đủ trong README.md và 12 file thiết kế trong thư mục Plan/.",
])

# ── 4. NHỮNG GÌ CHƯA LÀM / CÒN VẤN ĐỀ ─────────────────────
add_heading(doc, "4. Những vấn đề còn tồn tại và phương án xử lý", level=1)
add_para(
    doc,
    "Phần này liệt kê chi tiết các vấn đề được phát hiện khi đối chiếu với Debai.txt, "
    "giải thích từng vấn đề là gì, vì sao nó là vấn đề, và phương án khắc phục cụ thể."
)

# ── 4.1 ─────────────────────────────────────────────
add_heading(doc, "4.1 Mã chết và trùng lặp định nghĩa OrderCreatedEvent", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_para(
    doc,
    "Trong dự án hiện có TỚI BA cặp file mô tả cùng một nghiệp vụ \"sự kiện đã tạo đơn\" nhưng "
    "không nhất quán với nhau:"
)
add_bullets(doc, [
    "API/netcore/netcore.Commons/Contracts/Events/OrderCreatedEvent.cs – schema đầy đủ (OrderCode, UserId, TotalAmount, ProductName, UnitPrice, Quantity). KHÔNG được sử dụng ở bất kỳ luồng nào.",
    "API/netcore/netcore.Commons/Messages/Events/OrderCreatedEvent.cs – schema rút gọn (OrderCode, CustomerId, CreatedAt, Items chỉ có ProductId & Quantity). ĐANG được sử dụng thực tế.",
    "API/Services/API.Order/Messages/OrderCreatedMessage.cs – DTO trung gian map sang event \"Contracts\" – KHÔNG ai dùng.",
    "API/Services/API.Product/Consumers/OrderCreatedConsumer.cs – consumer dùng schema \"Contracts\", KHÔNG được đăng ký với MassTransit.",
    "API/Services/API.Product/Consumers/OrderCreatedEventConsumer.cs – consumer dùng schema \"Messages.Events\", ĐANG hoạt động.",
])
add_para(doc, "Đây là gì và vì sao quan trọng:", bold=True)
add_para(
    doc,
    "Event-driven là cốt lõi của giao tiếp bất đồng bộ trong microservice. Khi có hai định nghĩa event "
    "cùng tên nhưng schema khác nhau, người đọc code dễ nhầm lẫn về \"hợp đồng dữ liệu\" giữa các service. "
    "Ngoài ra, event đang dùng thực tế lại thiếu các trường quan trọng như ProductName, UnitPrice, "
    "TotalAmount – nếu sau này muốn thêm Notification Service gửi email xác nhận đơn hàng, hoặc thêm "
    "Reporting Service tổng hợp doanh thu theo event, sẽ phải sửa schema và rủi ro phá hợp đồng cũ."
)
add_para(doc, "Phương án khắc phục:", bold=True)
add_bullets(doc, [
    "Bước 1 – Chọn một schema chính thống: hợp nhất về netcore.Commons.Messages.Events.OrderCreatedEvent (đang chạy) và bổ sung các trường ProductName, UnitPrice, TotalAmount từ schema Contracts.",
    "Bước 2 – Xóa các file trùng lặp: netcore.Commons/Contracts/Events/OrderCreatedEvent.cs, API.Order/Messages/OrderCreatedMessage.cs, API.Product/Consumers/OrderCreatedConsumer.cs.",
    "Bước 3 – Cập nhật OrderService.cs để truyền đầy đủ thông tin sản phẩm vào event (đã có sẵn trong cartItems khi tạo đơn).",
    "Bước 4 – Build lại toàn bộ solution để chắc chắn không còn tham chiếu chết: dotnet build API/API.sln.",
])
add_para(
    doc,
    "Phương pháp kiểm chứng: sau khi sửa, đặt một đơn hàng test, kiểm tra log Seq xem event publish ra "
    "có đủ ProductName, UnitPrice không, đồng thời mở RabbitMQ UI (cổng 15672) kiểm tra payload thực tế "
    "trên queue order-created-queue.",
    italic=True,
)

# ── 4.2 ─────────────────────────────────────────────
add_heading(doc, "4.2 SQL Server không nằm trong docker-compose", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_para(
    doc,
    "File docker-compose.yml hiện chỉ orchestrate RabbitMQ, Seq, 4 service và Gateway. SQL Server "
    "phải được cài đặt và chạy độc lập trên máy host, sau đó người chạy phải chạy thủ công script "
    "API/database/sqlserver_full_setup.sql để tạo 4 database và seed dữ liệu."
)
add_para(doc, "Đây là gì và vì sao quan trọng:", bold=True)
add_para(
    doc,
    "Đề bài khuyến khích triển khai bằng Docker / Docker Compose để các service \"có thể chạy độc lập "
    "và phối hợp đúng\". Khi SQL Server đứng ngoài compose, người chấm/người mới phải tự cài SQL "
    "Server (Developer Edition rất nặng), tự cấu hình user sa, mật khẩu, chạy sqlcmd – làm tăng "
    "rào cản demo và dễ phát sinh lỗi không liên quan đến nghiệp vụ."
)
add_para(doc, "Phương án khắc phục:", bold=True)
add_bullets(doc, [
    "Thêm service mssql vào docker-compose.yml dùng image mcr.microsoft.com/mssql/server:2022-latest, biến môi trường SA_PASSWORD=viet123, ACCEPT_EULA=Y, mở cổng 1433.",
    "Thêm volume bind hoặc named volume cho /var/opt/mssql để dữ liệu không mất khi container restart.",
    "Tạo một service phụ db-init dùng image mcr.microsoft.com/mssql-tools chạy sqlcmd để nạp script sqlserver_full_setup.sql ngay sau khi mssql healthy (depends_on + condition: service_healthy).",
    "Đổi connection string trong appsettings.json từ host.docker.internal sang mssql (tên service) khi chạy bằng compose; vẫn cho phép override bằng biến môi trường ConnectionStrings__DefaultConnection để chạy ngoài Docker.",
])
add_para(
    doc,
    "Phương pháp kiểm chứng: từ máy hoàn toàn mới chỉ cần git clone + docker compose up --build phải "
    "chạy được tới mức đăng nhập admin/123456 thành công.",
    italic=True,
)

# ── 4.3 ─────────────────────────────────────────────
add_heading(doc, "4.3 Cấu hình Gateway còn rác và hard-code", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_bullets(doc, [
    "Key Gateway:ValidatePermissionEndpoint trong API/Gateway/FishShop.Gateway/appsettings.json:16 trỏ tới http://localhost:5001 – sẽ không truy cập được khi Gateway chạy trong container Docker (lúc đó phải là http://api-user:8080).",
    "Toàn bộ key ValidatePermissionEndpoint này KHÔNG được tham chiếu trong source code Gateway – tức là cấu hình rác.",
    "JwtSecret và ApiKey được commit thẳng vào appsettings.json.",
])
add_para(doc, "Đây là gì và vì sao quan trọng:", bold=True)
add_para(
    doc,
    "Cấu hình rác làm người đọc nhầm tưởng có một cơ chế kiểm tra quyền tập trung qua HTTP – nhưng "
    "thực tế Gateway chỉ verify JWT cục bộ. Việc hard-code localhost khiến cấu hình không di động "
    "giữa môi trường dev/Docker. Lộ secret trong git là thực hành xấu, dù trong đồ án học tập có "
    "thể chấp nhận."
)
add_para(doc, "Phương án khắc phục:", bold=True)
add_bullets(doc, [
    "Xóa key Gateway:ValidatePermissionEndpoint nếu không có kế hoạch sử dụng; hoặc triển khai middleware gọi tới endpoint này nếu thực sự cần kiểm tra quyền tập trung.",
    "Tách secret ra User Secrets (dotnet user-secrets) khi dev, và đọc qua biến môi trường khi chạy Docker; appsettings.json chỉ chứa placeholder.",
    "Trong docker-compose.yml thêm biến môi trường Gateway__JwtSecret và Gateway__ApiKey, đọc bằng configuration provider.",
    "Bổ sung file .env.example tài liệu hóa các biến môi trường cần thiết.",
])

# ── 4.4 ─────────────────────────────────────────────
add_heading(doc, "4.4 Thiếu cơ chế kiểm thử luồng nghiệp vụ liên service", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_para(
    doc,
    "Đề bài yêu cầu \"có cơ chế kiểm thử luồng nghiệp vụ liên service\". Hiện tại dự án chỉ có các file "
    ".http (REST Client) trong mỗi project service phục vụ test thủ công, không có test tự động (xUnit, "
    "integration test, hoặc kịch bản Postman có assertion)."
)
add_para(doc, "Đây là gì và vì sao quan trọng:", bold=True)
add_para(
    doc,
    "Trong hệ phân tán, một nghiệp vụ thường đi qua nhiều service (ví dụ: đăng nhập ở User Service → "
    "thêm vào giỏ ở Order Service → đặt hàng → publish event → trừ kho ở Product Service). Nếu không "
    "có kịch bản kiểm thử end-to-end, rất khó chứng minh khi bảo vệ rằng hệ thống hoạt động đúng "
    "khi các service phối hợp."
)
add_para(doc, "Phương án khắc phục:", bold=True)
add_bullets(doc, [
    "Tối thiểu (đủ để báo cáo): chuẩn hóa file .http hoặc viết một collection Postman có biến môi trường, gồm 6–8 request đại diện luồng \"đăng nhập → xem sản phẩm → thêm giỏ → đặt hàng → kiểm tra tồn kho giảm\". Kèm screenshot kết quả vào báo cáo.",
    "Mức tốt hơn: tạo project xUnit (FishShop.Tests) dùng WebApplicationFactory + Testcontainers (cho SQL Server và RabbitMQ) để chạy integration test thật, kiểm chứng OrderCreatedEvent thực sự đến và làm trừ kho.",
    "Mức nâng cao: viết kịch bản dotnet script hoặc shell/PowerShell tự động gọi API tuần tự và assert phản hồi – có thể nhúng vào CI sau này.",
])
add_para(
    doc,
    "Phương pháp kiểm chứng: chạy bộ test, chụp ảnh kết quả pass; với integration test, để dashboard "
    "Seq mở để cho người chấm thấy log luân chuyển giữa các service.",
    italic=True,
)

# ── 4.5 ─────────────────────────────────────────────
add_heading(doc, "4.5 Thiếu các thành phần nâng cao có thể ghi điểm cộng", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_para(
    doc,
    "Mục 5 của đề bài liệt kê các hạng mục nâng cao để cộng điểm, hiện tại đã có \"API Gateway\" và "
    "\"centralized logging\". Các hạng mục còn thiếu gồm:"
)
add_bullets(doc, [
    "Distributed tracing giữa các service (chưa có Correlation ID, chưa có OpenTelemetry).",
    "Circuit breaker / retry policy cho các HTTP call giữa service (chưa có Polly).",
    "Saga pattern hoặc cơ chế compensation khi consumer xử lý event thất bại quá số lần retry.",
    "Triển khai trên cloud hoặc Kubernetes.",
    "Dashboard giám sát service (Prometheus + Grafana hoặc tương đương).",
])
add_para(doc, "Phương án khắc phục theo độ ưu tiên (làm trước cái dễ, ăn điểm cao):", bold=True)
add_bullets(doc, [
    "(Ưu tiên 1 – dễ và đáng giá): Thêm OpenTelemetry vào tất cả service và Gateway, export sang Seq hoặc Jaeger. Cấu hình thuần code, chỉ cần thêm package OpenTelemetry.Extensions.Hosting + OpenTelemetry.Instrumentation.AspNetCore + Exporter, rồi bật trace propagation. Khi đặt 1 đơn hàng sẽ thấy trace xuyên Gateway → Order → RabbitMQ → Product.",
    "(Ưu tiên 2): Thêm Polly cho HttpClient ở những chỗ Order Service gọi sang Product Service (lấy thông tin sản phẩm). Cấu hình AddPolicyHandler retry 3 lần + circuit breaker mở khi 5 lỗi liên tiếp.",
    "(Ưu tiên 3): Cấu hình dead-letter queue cho consumer: trong MassTransit dùng cfg.UseInMemoryOutbox hoặc cấu hình _error queue tường minh để khi retry vượt giới hạn message rơi vào DLQ để xử lý sau.",
    "(Ưu tiên 4): Thêm pattern Saga đơn giản nhất – Order tạo ở trạng thái PENDING, chỉ chuyển sang CONFIRMED khi nhận được StockReservedEvent từ Product Service; nếu Product gửi StockFailedEvent thì Order chuyển sang CANCELLED. Đây là saga choreography, không cần thư viện ngoài.",
    "(Tùy chọn – tốn thời gian): Thêm Prometheus + Grafana hoặc dùng built-in dashboard của Seq để minh họa giám sát.",
])

# ── 4.6 ─────────────────────────────────────────────
add_heading(doc, "4.6 Báo cáo đồ án chưa đầy đủ theo cấu trúc yêu cầu", level=2)
add_para(doc, "Mô tả vấn đề:", bold=True)
add_para(
    doc,
    "Thư mục Plan/ hiện có 12 file thiết kế chi tiết (overview, kiến trúc, database, API contracts, "
    "frontend plan, conventions, deployment, ...). Tuy nhiên đây là tài liệu thiết kế tiền dự án, "
    "không phải báo cáo đồ án theo cấu trúc đề bài yêu cầu (Mục 6.2: mô tả bài toán, kiến trúc hệ "
    "thống, mô tả các microservice, thiết kế RESTful API, luồng message queue, thử nghiệm và đánh giá)."
)
add_para(doc, "Phương án khắc phục:", bold=True)
add_bullets(doc, [
    "Soạn báo cáo đồ án tổng hợp (file Word hoặc PDF) gồm các chương đúng cấu trúc đề bài yêu cầu.",
    "Tái sử dụng nội dung từ Plan/01_ARCHITECTURE.md, Plan/03_DATABASE.md, Plan/07_API_CONTRACTS.md để rút gọn vào báo cáo.",
    "Bổ sung chương \"Thử nghiệm và đánh giá\" gồm: kịch bản test, kết quả thực tế, hình chụp Seq/RabbitMQ UI, bảng so sánh trước/sau khi thêm message queue.",
    "Chèn sơ đồ kiến trúc (có thể dùng draw.io / Mermaid xuất PNG) và sơ đồ luồng event vào báo cáo cho trực quan.",
])

# ── 5. KẾT LUẬN ────────────────────────────────────────────
add_heading(doc, "5. Kết luận", level=1)
add_para(
    doc,
    "Dự án FishShop đã đáp ứng đầy đủ mức đạt tối thiểu của Debai.txt (Mục 7): có 4 microservice, "
    "RESTful API hoạt động đúng, có 1 luồng nghiệp vụ dùng message queue (tạo đơn hàng → trừ kho), "
    "các service chạy độc lập và phối hợp đúng. Hệ thống cũng đã có API Gateway, JWT, Docker Compose, "
    "centralized logging với Seq – là các điểm khuyến khích ở Mục 4."
)
add_para(
    doc,
    "Tuy nhiên vẫn còn các vấn đề cần xử lý để báo cáo đầy đủ và để hệ thống chuyên nghiệp hơn: "
    "(1) dọn mã chết quanh OrderCreatedEvent, (2) đưa SQL Server vào docker-compose, (3) làm sạch "
    "cấu hình Gateway, (4) bổ sung kịch bản kiểm thử liên service, (5) soạn báo cáo đồ án theo cấu "
    "trúc Mục 6 của đề bài."
)
add_para(
    doc,
    "Nếu thời gian cho phép, nên triển khai thêm OpenTelemetry và Polly để ghi điểm phần nâng cao. "
    "Các phương pháp khắc phục đã được nêu rõ ở phần 4 kèm bước thực hiện và cách kiểm chứng."
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end.add_run("— Hết báo cáo —")
end_run.italic = True

output_path = r"C:\Users\Admin\Desktop\HTPT\HTPT\report\BaoCao_DanhGia_DoAn.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
